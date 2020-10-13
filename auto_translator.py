from urllib import request, parse
import json

import requests
import time
import random
import urllib.request
import urllib.parse
import json
from docx import Document
from docx.shared import Inches

import re


input_file_path = "2captions-FedericoTemporal.docx"
output_file_path = "trans4cuocuo-" + input_file_path

def youdao_translate(content):
    ##### test #####
    # print("翻译原文:%s" % content)
    # return content
    
    ##### test1 #####
    # print("翻译原文:%s" % content)
    # target = content
    # url = 'http://fanyi.youdao.com/translate?smartresult=dict&smartresult=rule'
    # data = {
        # 'i' : target,
        # 'from' : 'AUTO',
        # 'to' : 'AUTO',
        # 'smartresult' : 'dict',
        # 'client' : 'fanyideskweb',
        # # 'salt' : '15810537039389',
        # # 'sign' : '157b38258a2253c7899895880487edfd',
        # # 'ts' : '1581053703938',
        # # 'bv' : '901200199a98c590144a961dac532964',
        
        # 'salt' : '16002467618444',
        # 'sign' : 'e61109f1cab5ece346bfafc019944f04',
        # 'lts' : '1600246761844',
        # 'bv' : 'e2a78ed30c66e16a857c5b6486a1d326',
        
        # # 'salt' : '16002470619810',
        # # 'sign' : '9a9bae39d5f9f80c48ffeafe1d4694cf',
        # # 'lts' : '1600247061981',
        # # 'bv' : 'e2a78ed30c66e16a857c5b6486a1d326',
        
        # 'doctype' : 'json',
        # 'version' : '2.1',
        # 'keyfrom' : 'fanyi.web',
        # # 'action' : 'FY_BY_CLICKBUTTION',    
        # 'action' : 'FY_BY_REALTlME'
        
    # }

    # head = {'User-Agent' : 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/79.0.3945.130 Safari/537.36'}

    # data = urllib.parse.urlencode(data).encode('utf-8')

    # rep = urllib.request.Request(url, data, head)
    # response = urllib.request.urlopen(rep)

    # html = response.read().decode('utf-8')
    # result = json.loads(html)
    # result = result['translateResult'][0][0]['tgt']

    # print("翻译结果为:",result)
    # return result




    ##### test2 #####
    print("翻译原文:%s" % content)
    req_url = 'http://fanyi.youdao.com/translate'  # 创建连接接口
    # 创建要提交的数据
    Form_Date = {}
    Form_Date['i'] = content
    Form_Date['doctype'] = 'json'
    Form_Date['form'] = 'AUTO'
    Form_Date['to'] = 'AUTO'
    Form_Date['smartresult'] = 'dict'
    Form_Date['client'] = 'fanyideskweb'
    Form_Date['salt'] = '1526995097962'
    Form_Date['sign'] = '8e4c4765b52229e1f3ad2e633af89c76'
    Form_Date['version'] = '2.1'
    Form_Date['keyform'] = 'fanyi.web'
    Form_Date['action'] = 'FY_BY_REALTIME'
    Form_Date['typoResult'] = 'false'
 
    data = parse.urlencode(Form_Date).encode('utf-8') #数据转换
    response = request.urlopen(req_url, data) #提交数据并解析
    html = response.read().decode('utf-8')  #服务器返回结果读取
    print(html)
    # 可以看出html是一个json格式
    translate_results = json.loads(html)  #以json格式载入
    translate_results = translate_results['translateResult'][0][0]['tgt']  # json格式调取
    print("翻译结果：%s" % translate_results)
    return translate_results #输出结果.

 
    

if __name__ == '__main__':

    # youdao_translate("are you ok?")
    document = Document(input_file_path)  #打开文件demo.docx
    document_write = Document()
    # 依次处理每一句话，在当前句号处输出
    count = 0
    translate_buffer = []
    sentence = ''
    for paragraph in document.paragraphs:
        document_write.add_paragraph(paragraph.text)        
        count += 1
        print("count:%d" % count)
        if count != 1 and count != 2: # 数字/时间 跳过
            content = paragraph.text
            print("content:%s" % content)
            print("sentence:%s" % sentence)
            # input()
            # sen_array = content.split('.') # 分割成N个句子，其中前后句子不完整，中间句子完整
            sen_array = re.split('[.:?]', content)
            
            for i in sen_array:
                print("sen_array:%s" % i)
                
            for i in range(len(sen_array)):
                if len(sen_array) == 1: # 没有句号
                    print("log: 没有句号: %s" % sen_array[0])
                    print("sentence:%s" % sentence)
                    sentence = sentence + " " + sen_array[0]
                    print("sentence:%s" % sentence)
                else:
                    print("log: 有句号: %s" % content)
                    if i != 0 and i != len(sen_array)-1: # 翻译并输出   
                        print("i:%d" % i)
                        print("log: mid part: %s" % sen_array[i])
                        # sentence 送去翻译
                        result = youdao_translate(i)
                        # 得到翻译结果，填入word
                        translate_buffer.append(result)
                    elif i == 0: 
                        print("log: first part: %s" % sen_array[i])
                        # 链接到前一个句子，翻译并输出
                        sentence = sentence + " " + sen_array[i]
                        # sentence 送去翻译
                        result = youdao_translate(sentence)
                        # 得到翻译结果，填入word
                        translate_buffer.append(result)
                    else:
                        print("log: last part: %s" % sen_array[i])
                        # 记录句子，保留到下一阶段输出
                        sentence = sen_array[i]
            if paragraph.text == '' or paragraph.text == ' ':
                for i in translate_buffer:
                    print("输出到doc：%s" % i)
                    # input()
                    #document_write.add_paragraph("译文："+str(i))
                    document_write.add_paragraph(str(i))
                    document_write.add_paragraph(" ")
                translate_buffer = []
                count = 0
    document_write.save(output_file_path)  

    
    
